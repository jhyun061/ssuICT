# app.py
# =========================================================
# AI 자기소개서 코칭 - 카카오톡 스타일 UI
# =========================================================
# 설치: pip install streamlit python-docx reportlab langchain langchain-openai python-dotenv
# 실행: streamlit run app.py
# =========================================================

import os, io, datetime, json
from typing import Optional, List, Dict
import streamlit as st

# ===== 문서 생성 라이브러리 (선택) =====
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    DOC_LIBS_AVAILABLE = True
except:
    DOC_LIBS_AVAILABLE = False

# ===== LangChain (선택) =====
try:
    from langchain_openai import ChatOpenAI
    from langchain.prompts import ChatPromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except:
    LANGCHAIN_AVAILABLE = False

# ================= 페이지 설정 =================
st.set_page_config(
    page_title="AI 자기소개서 코칭",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ================= 카카오톡 스타일 CSS =================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;600;700&display=swap');
    
    /* 전체 배경 및 기본 스타일 */
    .stApp {
        background: #b2c7d9;
        font-family: 'Noto Sans KR', sans-serif;
    }
    
    /* 메인 컨테이너 */
    .main .block-container {
        padding: 0;
        max-width: 100%;
        margin: 0;
    }
    
    /* 기본 탭 숨김 */
    .stTabs [data-baseweb="tab-list"] {
        display: none;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        padding: 0;
    }
    
    /* 상단 헤더 */
    .chat-header {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        height: 60px;
        background: rgba(0, 0, 0, 0.85);
        color: white;
        display: flex;
        align-items: center;
        justify-content: center;
        z-index: 1000;
        backdrop-filter: blur(10px);
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .header-title {
        font-size: 18px;
        font-weight: 600;
        color: white;
    }
    
    /* 채팅 영역 */
    .chat-container {
        margin-top: 60px;
        margin-bottom: 120px;
        padding: 20px;
        min-height: calc(100vh - 180px);
        overflow-y: auto;
    }
    
    /* 메시지 버블 */
    .msg-row {
        display: flex;
        margin-bottom: 15px;
        align-items: flex-end;
    }
    
    .msg-row.user {
        justify-content: flex-end;
    }
    
    .msg-row.ai {
        justify-content: flex-start;
    }
    
    .msg-bubble {
        max-width: 70%;
        padding: 10px 14px;
        border-radius: 18px;
        font-size: 14px;
        line-height: 1.5;
        word-break: break-word;
        box-shadow: 0 1px 2px rgba(0, 0, 0, 0.15);
        position: relative;
    }
    
    .msg-bubble.user {
        background: #ffeb33;
        color: #000;
        border-top-right-radius: 4px;
    }
    
    .msg-bubble.ai {
        background: white;
        color: #000;
        border-top-left-radius: 4px;
    }
    
    .msg-time {
        font-size: 11px;
        color: #888;
        margin: 0 8px;
        white-space: nowrap;
    }
    
    /* 설정 페이지 */
    .settings-container {
        margin-top: 60px;
        margin-bottom: 60px;
        padding: 20px;
        background: white;
        min-height: calc(100vh - 120px);
    }
    
    .settings-section {
        background: #f8f9fa;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 15px;
        border: 1px solid #e9ecef;
    }
    
    .settings-title {
        font-size: 16px;
        font-weight: 600;
        margin-bottom: 15px;
        color: #333;
    }
    
    /* 저장소 페이지 */
    .storage-container {
        margin-top: 60px;
        margin-bottom: 60px;
        padding: 20px;
        background: white;
        min-height: calc(100vh - 120px);
    }
    
    .file-item {
        background: #f8f9fa;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 10px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        border: 1px solid #e9ecef;
    }
    
    .file-info {
        flex: 1;
    }
    
    .file-name {
        font-weight: 500;
        margin-bottom: 5px;
    }
    
    .file-date {
        font-size: 12px;
        color: #888;
    }
    
    /* 버튼 스타일 */
    .stButton > button {
        background: #ffeb33;
        color: #000;
        border: none;
        border-radius: 20px;
        padding: 8px 20px;
        font-weight: 500;
        transition: all 0.2s;
        width: 100%;
    }
    
    .stButton > button:hover {
        background: #ffd900;
        color: #000;
    }
    
    /* 입력창 스타일 */
    .stTextInput > div > div > input {
        background: #f5f5f5;
        border: 1px solid #e0e0e0;
        border-radius: 20px;
        padding: 10px 15px;
        font-size: 14px;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #ffeb33 !important;
        box-shadow: 0 0 0 2px rgba(255, 235, 51, 0.2) !important;
    }
    
    /* 파일 업로드 스타일 */
    .stFileUploader > label {
        background: #f8f9fa;
        border: 2px dashed #dee2e6;
        border-radius: 10px;
        padding: 20px;
        text-align: center;
    }
    
    /* 스크롤바 */
    ::-webkit-scrollbar {
        width: 6px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #888;
        border-radius: 3px;
    }
    
    /* selectbox 스타일 */
    .stSelectbox > div > div {
        background: #f5f5f5;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
    }
    
    /* 슬라이더 스타일 */
    .stSlider > div > div > div {
        color: #ffeb33;
    }
    
    /* 정보 박스 스타일 */
    .stInfo {
        background: #e3f2fd;
        border: 1px solid #bbdefb;
        border-radius: 8px;
    }
    
    /* 성공 메시지 스타일 */
    .stSuccess {
        background: #e8f5e8;
        border: 1px solid #c8e6c9;
        border-radius: 8px;
    }
    
    /* 경고 메시지 스타일 */
    .stWarning {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 8px;
    }
    
    /* 채팅 입력 영역 고정 */
    .chat-input-section {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background: white;
        padding: 15px;
        border-top: 1px solid #e0e0e0;
        z-index: 999;
    }
    
    /* 빠른 답변 버튼 */
    .quick-replies {
        margin-bottom: 10px;
    }
    
    .quick-reply-btn {
        display: inline-block;
        padding: 6px 12px;
        margin: 2px;
        background: white;
        border: 1px solid #e0e0e0;
        border-radius: 15px;
        font-size: 12px;
        cursor: pointer;
        transition: all 0.2s;
    }
    
    .quick-reply-btn:hover {
        background: #ffeb33;
        border-color: #ffeb33;
    }
    
    /* 반응형 디자인 */
    @media (max-width: 768px) {
        .msg-bubble {
            max-width: 85%;
        }
        
        .settings-container, .storage-container {
            padding: 10px;
        }
        
        .chat-container {
            padding: 10px;
        }
    }
</style>
""", unsafe_allow_html=True)

# ================= 세션 초기화 =================
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({
        "role": "ai",
        "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
        "time": datetime.datetime.now().strftime("%H:%M")
    })

if "current_tab" not in st.session_state:
    st.session_state.current_tab = "대화"

if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")

if "saved_files" not in st.session_state:
    st.session_state.saved_files = []

if "save_format" not in st.session_state:
    st.session_state.save_format = "txt"

if "model_settings" not in st.session_state:
    st.session_state.model_settings = {
        "temperature": 0.7,
        "max_length": 1000,
        "tone": "professional"
    }

# ================= 가이드라인 응답 =================
def get_guideline_response():
    return """📝 **AI 자기소개서 입력 가이드라인**

**1. 구체적으로 질문하기**
✅ "마케팅 직무 신입 자기소개서 도입부 작성해줘"
❌ "자소서 써줘"

**2. 배경 정보 제공하기**
• 지원 회사와 직무
• 본인의 주요 경험
• 강조하고 싶은 역량

**3. 효과적인 질문 예시**
• "고객 서비스 경험을 영업직무에 연결하는 방법"
• "프로젝트 경험을 STAR 기법으로 정리해줘"
• "IT 기업 지원동기 작성 도와줘"

**4. 첨삭 요청 방법**
• 작성한 내용 복사 후 "이 내용 첨삭해줘"
• 파일 업로드 후 "구체성 높여줘"
• "이 문장 더 임팩트 있게 수정해줘"

**5. 단계별 접근**
1️⃣ 전체 구조 잡기
2️⃣ 각 문단 작성
3️⃣ 표현 다듬기
4️⃣ 최종 검토

💡 **Tip**: 한 번에 모든 걸 해결하려 하지 말고, 단계별로 질문하세요!"""

# ================= AI 응답 생성 =================
def get_ai_response(user_input: str, uploaded_file=None) -> str:
    # 가이드라인 요청 체크
    guideline_keywords = ["가이드", "가이드라인", "도움말", "사용법", "어떻게"]
    if any(keyword in user_input for keyword in guideline_keywords):
        return get_guideline_response()
    
    # 템플릿 응답 (API 키 없을 때)
    if not st.session_state.api_key or not LANGCHAIN_AVAILABLE:
        templates = {
            "default": """자기소개서 작성을 도와드리겠습니다!

구체적으로 알려주시면 더 정확한 도움을 드릴 수 있어요:
• 어떤 직무에 지원하시나요?
• 어떤 부분이 어려우신가요?
• 특별히 강조하고 싶은 경험이 있나요?""",
            
            "첨삭": """자기소개서 첨삭 포인트를 알려드릴게요:

✅ 구체적인 숫자와 성과 포함
✅ 직무와 연관된 경험 강조
✅ 문장은 간결하고 명확하게
✅ 진정성 있는 지원동기

파일을 업로드하거나 내용을 보내주시면 더 자세히 봐드릴게요!""",
            
            "시작": """자기소개서 작성을 시작해볼까요?

**Step 1. 기본 정보**
• 지원 회사:
• 지원 직무:
• 경력 구분: (신입/경력)

이 정보를 알려주시면 맞춤형으로 도와드릴게요!"""
        }
        
        if "첨삭" in user_input or "수정" in user_input:
            return templates["첨삭"]
        elif "시작" in user_input or "처음" in user_input:
            return templates["시작"]
        else:
            return templates["default"]
    
    # LangChain을 이용한 AI 응답
    try:
        llm = ChatOpenAI(
            api_key=st.session_state.api_key,
            model="gpt-4o-mini",
            temperature=st.session_state.model_settings["temperature"]
        )
        
        system_prompt = f"""당신은 전문 자기소개서 작성 코치입니다.
        톤: {st.session_state.model_settings["tone"]}
        최대 길이: {st.session_state.model_settings["max_length"]}자
        
        - 구체적이고 실용적인 조언
        - 예시를 들어 설명
        - 친근하면서도 전문적인 톤
        - 이모지는 최소한으로 사용"""
        
        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOC_LIBS_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = "파일을 읽을 수 없습니다."
                
                user_input = f"다음 자기소개서를 검토하고 개선점을 제안해주세요:\n\n{content}\n\n{user_input}"
            except Exception as e:
                return f"파일 처리 중 오류: {e}"
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        
        chain = LLMChain(llm=llm, prompt=prompt)
        response = chain.invoke({"input": user_input})
        
        return response.get("text", str(response))
        
    except Exception as e:
        return f"오류가 발생했습니다. 다시 시도해주세요.\n{str(e)}"

# ================= 대화 저장 =================
def save_conversation():
    content = ""
    for msg in st.session_state.messages:
        role = "👤 사용자" if msg["role"] == "user" else "🤖 AI 코치"
        content += f"[{msg.get('time', '')}] {role}\n{msg['content']}\n\n"
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"자소서대화_{timestamp}"
    
    # 선택된 형식으로 저장
    if st.session_state.save_format == "txt":
        file_data = content
        mime = "text/plain"
        ext = "txt"
    elif st.session_state.save_format == "docx" and DOC_LIBS_AVAILABLE:
        doc = Document()
        doc.add_heading('AI 자기소개서 코칭 대화', 0)
        for para in content.split('\n'):
            doc.add_paragraph(para)
        
        bio = io.BytesIO()
        doc.save(bio)
        file_data = bio.getvalue()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        file_data = content
        mime = "text/plain"
        ext = "txt"
    
    # 저장 목록에 추가
    st.session_state.saved_files.append({
        "name": f"{filename}.{ext}",
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "size": len(file_data),
        "data": file_data,
        "mime": mime
    })
    
    return f"{filename}.{ext}"

# ================= UI 렌더링 함수 =================
def render_header():
    st.markdown(f'''
        <div class="chat-header">
            <div class="header-title">AI 자기소개서 코칭</div>
        </div>
    ''', unsafe_allow_html=True)

def render_chat_tab():
    # 채팅 메시지 표시
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f'''
                <div class="msg-row user">
                    <div class="msg-time">{msg.get("time", "")}</div>
                    <div class="msg-bubble user">{msg["content"]}</div>
                </div>
            ''', unsafe_allow_html=True)
        else:
            content_html = msg["content"].replace('\n', '<br>')
            st.markdown(f'''
                <div class="msg-row ai">
                    <div class="msg-bubble ai">{content_html}</div>
                    <div class="msg-time">{msg.get("time", "")}</div>
                </div>
            ''', unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # 입력 영역 (고정)
    with st.container():
        st.markdown('<div class="chat-input-section">', unsafe_allow_html=True)
        
        # 빠른 답변
        st.markdown('<div class="quick-replies">', unsafe_allow_html=True)
        quick_replies = ["가이드라인 알려줘", "자소서 시작하기", "첨삭 받고 싶어", "예시 보여줘"]
        cols = st.columns(len(quick_replies))
        for i, reply in enumerate(quick_replies):
            with cols[i]:
                if st.button(reply, key=f"quick_{i}"):
                    st.session_state.messages.append({
                        "role": "user",
                        "content": reply,
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                    response = get_ai_response(reply)
                    st.session_state.messages.append({
                        "role": "ai",
                        "content": response,
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
        
        # 파일 업로드
        uploaded_file = st.file_uploader(
            "📎 파일 첨부 (txt, docx)",
            type=['txt', 'docx'],
            label_visibility="visible",
            help="자기소개서 파일을 업로드하여 첨삭받을 수 있습니다."
        )
        
        # 메시지 입력
        with st.form("chat_form", clear_on_submit=True):
            col1, col2 = st.columns([5, 1])
            with col1:
                user_input = st.text_input(
                    "메시지",
                    placeholder="메시지를 입력하세요...",
                    label_visibility="collapsed"
                )
            with col2:
                send = st.form_submit_button("전송")
            
            if send and user_input:
                # 사용자 메시지 추가
                st.session_state.messages.append({
                    "role": "user",
                    "content": user_input,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                
                # AI 응답 생성
                with st.spinner("입력 중..."):
                    response = get_ai_response(user_input, uploaded_file)
                
                st.session_state.messages.append({
                    "role": "ai",
                    "content": response,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                
                st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

def render_settings_tab():
    st.markdown('<div class="settings-container">', unsafe_allow_html=True)
    
    st.title("⚙️ 설정")
    
    # API 설정
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">🔑 API 설정</div>', unsafe_allow_html=True)
    
    api_key = st.text_input(
        "OpenAI API Key",
        value=st.session_state.api_key,
        type="password",
        placeholder="sk-...",
        help="OpenAI API 키를 입력하면 더 정확한 AI 응답을 받을 수 있습니다."
    )
    
    if api_key != st.session_state.api_key:
        st.session_state.api_key = api_key
        st.success("API 키가 저장되었습니다!")
    
    st.info("💡 API 키가 없어도 기본 기능을 사용할 수 있습니다")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # 대화 관리
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">💬 대화 관리</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ 대화 초기화"):
            st.session_state.messages = [{
                "role": "ai",
                "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
                "time": datetime.datetime.now().strftime("%H:%M")
            }]
            st.success("대화가 초기화되었습니다!")
            st.rerun()
    
    with col2:
        if st.button("💾 대화 저장"):
            filename = save_conversation()
            st.success(f"{filename} 저장됨!")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

def render_advanced_settings_tab():
    st.markdown('<div class="settings-container">', unsafe_allow_html=True)
    
    st.title("🔧 세부설정")
    
    # AI 모델 설정
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">🤖 AI 모델 설정</div>', unsafe_allow_html=True)
    
    st.session_state.model_settings["temperature"] = st.slider(
        "창의성 (Temperature)",
        min_value=0.0,
        max_value=1.0,
        value=st.session_state.model_settings["temperature"],
        step=0.1,
        help="값이 높을수록 더 창의적인 답변을 생성합니다."
    )
    
    st.session_state.model_settings["max_length"] = st.number_input(
        "최대 응답 길이 (자)",
        min_value=100,
        max_value=3000,
        value=st.session_state.model_settings["max_length"],
        step=100,
        help="AI 응답의 최대 길이를 설정합니다."
    )
    
    st.session_state.model_settings["tone"] = st.selectbox(
        "응답 톤",
        ["professional", "friendly", "casual", "formal"],
        index=["professional", "friendly", "casual", "formal"].index(st.session_state.model_settings["tone"]),
        help="AI의 응답 스타일을 선택합니다."
    )
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # 저장 설정
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">💾 저장 설정</div>', unsafe_allow_html=True)
    
    st.session_state.save_format = st.selectbox(
        "기본 저장 형식",
        ["txt", "docx", "pdf"],
        index=["txt", "docx", "pdf"].index(st.session_state.save_format),
        help="대화 저장 시 사용할 파일 형식을 선택합니다."
    )
    
    st.info("📌 저장된 파일은 '저장소' 탭에서 확인할 수 있습니다")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

def render_storage_tab():
    st.markdown('<div class="storage-container">', unsafe_allow_html=True)
    
    st.title("📁 저장소")
    
    if not st.session_state.saved_files:
        st.info("저장된 파일이 없습니다. 대화를 저장하려면 설정 탭을 이용하세요.")
    else:
        st.write(f"총 {len(st.session_state.saved_files)}개의 파일이 저장되어 있습니다.")
        
        for i, file in enumerate(st.session_state.saved_files):
            with st.container():
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.markdown(f'''
                        <div class="file-item">
                            <div class="file-info">
                                <div class="file-name">📄 {file["name"]}</div>
                                <div class="file-date">{file["date"]} · {file["size"]} bytes</div>
                            </div>
                        </div>
                    ''', unsafe_allow_html=True)
                
                with col2:
                    st.download_button(
                        label="다운로드",
                        data=file["data"],
                        file_name=file["name"],
                        mime=file["mime"],
                        key=f"download_{i}_{file['name']}"
                    )
    
    # 일괄 삭제
    if st.session_state.saved_files:
        st.markdown("---")
        if st.button("🗑️ 모든 파일 삭제"):
            st.session_state.saved_files = []
            st.success("모든 파일이 삭제되었습니다!")
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# ================= 메인 앱 =================
def main():
    # 헤더
    render_header()
    
    # 탭 생성 (Streamlit 내장 탭 사용)
    tab1, tab2, tab3, tab4 = st.tabs(["💬 대화", "⚙️ 설정", "🔧 세부설정", "📁 저장소"])
    
    with tab1:
        render_chat_tab()
    
    with tab2:
        render_settings_tab()
    
    with tab3:
        render_advanced_settings_tab()
    
    with tab4:
        render_storage_tab()

# 프로그램 진입점
if __name__ == "__main__":
    main()
